"""
Citation Verification Tool for AURA Policy Documents.

Pipeline:
  1. Extract citations (URL/DOI + surrounding claim) from PDF or DOCX
  2. Resolve sources  — DOI → CrossRef → Unpaywall (open-access URL)
  3. Verify quotes    — fuzzy-match claim against source full text
  4. Screenshot       — navigate to source, highlight matched passage, capture
  5. Report           — annotated PDF with screenshot log + summary

Entry point:
    result = await run_verification("/abs/path/to/doc.pdf")
    # result.pdf_path → path to generated report PDF
"""

import asyncio
import io
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Callable, Optional

from PIL import Image, ImageDraw, ImageFont

# Open-access resolution functions live in their own MCP tool — imported here
# so citation_verifier can use them internally without duplication.
from app.tools.open_access_resolver import (
    resolve_doi      as _resolve_doi,
    core_fetch_by_doi as _core_fetch_by_doi,
    extract_doi_from_url as _extract_doi_from_page,
    fetch_pdf_text   as _fetch_pdf_text,
    fetch_via_archives as _fetch_via_archives,
)

logger = logging.getLogger(__name__)



# ── Score thresholds ──────────────────────────────────────────────────────────
_SCORE_CONFIRMED  = 0.75
_SCORE_PARTIAL    = 0.45
_SCORE_UNCERTAIN  = 0.20

_BROWSER_TIMEOUT  = 20_000  # ms Playwright timeout

# Maximum source characters to fuzzy-search over
_MAX_SOURCE_CHARS = 50_000


# ── Data models ───────────────────────────────────────────────────────────────

@dataclass
class Citation:
    index: int
    claim: str            # Sentence(s) from the document that include this citation
    raw_ref: str          # Raw URL/DOI as found in the doc
    source_url: str       # URL to fetch (may be same as raw_ref if already a URL)
    doi: Optional[str] = None

    # Populated after resolution
    resolved_url: Optional[str] = None
    page_title: Optional[str] = None
    source_text: Optional[str] = None

    # Populated after verification
    match_score: float = 0.0
    matched_passage: Optional[str] = None
    status: str = "pending"   # confirmed | partial | uncertain | hallucinated | unreachable

    # Populated after screenshot + annotation
    screenshot_png: Optional[bytes] = None
    annotated_png: Optional[bytes] = None

    claim_source: str = "body"      # "body" | "ref_only"
    highlight_found: bool = False   # True if the quote was located on the page

    # Populated after relational entity back-check (local workhorse)
    llm_verdict: Optional[str] = None   # CONFIRMED | PARTIAL | NOT_SUPPORTED
    llm_note:    Optional[str] = None   # Entity-level explanation from workhorse

    error: Optional[str] = None


@dataclass
class VerificationReport:
    doc_path: str
    doc_name: str
    citations: list[Citation] = field(default_factory=list)
    pdf_path: Optional[str] = None
    generated_at: str = field(default_factory=lambda: datetime.now().isoformat())
    doc_summary: str = ""           # Local LLM-generated overview for PDF cover

    @property
    def summary(self) -> dict:
        counts = {
            "confirmed": 0, "partial": 0, "uncertain": 0,
            "hallucinated": 0, "unreachable": 0,
        }
        for c in self.citations:
            if c.status in counts:
                counts[c.status] += 1
        return {"total": len(self.citations), **counts}


# ── Step 1: Citation Extraction ───────────────────────────────────────────────

_URL_RE  = re.compile(r'https?://[^\s\]\)\"\'\,\;\:>]+', re.IGNORECASE)
_DOI_RE  = re.compile(r'\b(10\.\d{4,9}/[^\s\]\)\"\'\,\;\:>]+)', re.IGNORECASE)
_DOI_LABEL_RE = re.compile(r'doi:\s*(10\.\d{4,9}/[^\s\]\)\"\'\,\;\:>]+)', re.IGNORECASE)


def _split_sentences(text: str) -> list[str]:
    """Naive but reliable sentence splitter for policy documents."""
    return re.split(r'(?<=[.!?])\s{1,2}(?=[A-Z\"\'])', text)


_BULLET_RE   = re.compile(r'^[\s\u2022\u2023\u25e6\u2043\u2219\ufffd\t•‣◦⁃●\-]+')
_FOOTNOTE_RE = re.compile(r'\d+(?:,\s*\d+)*\s*$')   # trailing footnote numbers


def _build_claim(sentences: list[str], cite_idx: int, window: int = 2) -> str:
    """Return up to `window` sentences before + the sentence containing the citation."""
    start = max(0, cite_idx - window)
    raw = " ".join(sentences[start:cite_idx + 1])
    # Strip refs, bullet chars, and trailing footnote markers so claim is clean prose
    raw = _URL_RE.sub("", raw)
    raw = _DOI_RE.sub("", raw)
    raw = _BULLET_RE.sub("", raw)
    raw = _FOOTNOTE_RE.sub("", raw)
    return " ".join(raw.split()).strip()


_REF_HEADER_RE = re.compile(
    r'^\s*(references?|works cited|bibliography|endnotes?|notes)\s*$', re.I
)
_REF_ITEM_RE  = re.compile(r'^\s*([\d,\s]+?)[\s\.,]')
_FOOT_RE      = re.compile(r'\.?([\d]+(?:,\s*\d+)*)\s*$')


def _split_body_refs_pdf(path: str) -> tuple[str, str]:
    """Return (body_text, ref_text) split at the References section."""
    from pypdf import PdfReader
    full = "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
    m = re.search(r'\n\s*(References?|Works Cited|Bibliography)\s*\n', full, re.I)
    if m:
        return full[:m.start()].strip(), full[m.start():].strip()
    return full, ""


def _split_body_refs_docx(path: str) -> tuple[str, str]:
    """Return (body_text, ref_text) split at the References section."""
    from docx import Document
    body, refs = [], []
    in_refs = False
    for para in Document(path).paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if not in_refs and _REF_HEADER_RE.match(text):
            in_refs = True
            continue
        (refs if in_refs else body).append(text)
    return "\n".join(body), "\n".join(refs)


def _parse_ref_list(ref_text: str) -> dict[int, str]:
    """
    Parse a numbered reference list into {ref_number: url}.
    Handles multi-number entries like '3,4 Oklahoma Energy Today...'.
    """
    result: dict[int, str] = {}
    # Split at lines that start with digits
    entries = re.split(r'\n(?=\s*[\d])', ref_text)
    for entry in entries:
        m = _REF_ITEM_RE.match(entry)
        if not m:
            continue
        nums = [int(n.strip()) for n in re.split(r'[,\s]+', m.group(1).strip())
                if n.strip().isdigit()]
        if not nums:
            continue
        urls = _URL_RE.findall(entry)
        dois = _DOI_RE.findall(entry)
        url = None
        if urls:
            url = urls[-1].rstrip(".,;)")   # last URL = actual source
        elif dois:
            url = f"https://doi.org/{dois[0].rstrip('.,;)')}"
        if url:
            for num in nums:
                result[num] = url
    return result


def _find_claim_for_ref(sentences: list[str], ref_num: int) -> str:
    """
    Find the body sentence(s) that bear footnote marker `ref_num`.
    Markers appear as bare digits at the end of a sentence (e.g. '...approvals.1'),
    or mid-sentence when the text continues after the citation.
    """
    # Primary: marker at end of line (most common)
    for idx, sentence in enumerate(sentences):
        m = _FOOT_RE.search(sentence.rstrip())
        if m:
            nums = {int(n.strip()) for n in m.group(1).split(",") if n.strip().isdigit()}
            if ref_num in nums:
                return _build_claim(sentences, idx)

    # Fallback: marker anywhere in line as a standalone number (mid-sentence footnotes)
    pattern = re.compile(r'(?<!\d)' + str(ref_num) + r'(?!\d)')
    for idx, sentence in enumerate(sentences):
        if pattern.search(sentence):
            return _build_claim(sentences, idx)

    return ""


def _build_docx_ref_claims(doc_path: str, ref_map: dict[int, str]) -> dict[int, str]:
    """
    Build {ref_num: claim_text} directly from DOCX superscript runs.
    More reliable than text regex — handles both end-of-line and mid-sentence markers.
    """
    from docx import Document
    doc = Document(doc_path)

    # Find where the reference section starts so we only look at body paragraphs
    ref_start = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs):
        if _REF_HEADER_RE.match(para.text.strip()):
            ref_start = i
            break

    # Build para_index → list[ref_nums] from superscript runs
    para_to_refs: dict[int, list[int]] = {}
    for i, para in enumerate(doc.paragraphs[:ref_start]):
        for run in para.runs:
            if run.font.superscript:
                # Superscript may be "29" or "3,4" etc
                for tok in re.split(r'[,\s]+', run.text.strip()):
                    if tok.isdigit() and int(tok) in ref_map:
                        para_to_refs.setdefault(i, []).append(int(tok))

    # Build body lines for context windowing
    body_lines = [p.text.strip() for p in doc.paragraphs[:ref_start] if p.text.strip()]

    ref_to_claim: dict[int, str] = {}
    for para_idx, ref_nums in para_to_refs.items():
        # Rebuild relative index in body_lines
        line_text = doc.paragraphs[para_idx].text.strip()
        try:
            line_idx = body_lines.index(line_text)
        except ValueError:
            line_idx = 0
        claim = _build_claim(body_lines, line_idx, window=0)
        for rn in ref_nums:
            if rn not in ref_to_claim:
                ref_to_claim[rn] = claim

    return ref_to_claim


def extract_citations(doc_path: str) -> list[Citation]:
    """
    Parse a PDF or DOCX and return all (claim, source_url) Citation objects.

    Two-pass strategy for numbered reference lists (common in policy documents):
      Pass 1 — scan the reference section to collect {ref_num: url}
      Pass 2 — find each ref number's footnote marker in body text; build
                claim from the surrounding body sentences only.

    Falls back to inline URL/DOI scan (original behaviour) when no numbered
    reference list is detected.
    """
    suffix = Path(doc_path).suffix.lower()
    if suffix == ".pdf":
        body_text, ref_text = _split_body_refs_pdf(doc_path)
    elif suffix in (".docx", ".doc"):
        body_text, ref_text = _split_body_refs_docx(doc_path)
    else:
        raise ValueError(f"Unsupported document type: {suffix!r} — expected .pdf or .docx")

    citations: list[Citation] = []
    seen: set[str] = set()

    # ── Path A: numbered reference list ──────────────────────────────────────
    if ref_text:
        ref_map = _parse_ref_list(ref_text)   # {num: url}
        if ref_map:
            # For DOCX: use superscript run metadata to map refs → claims precisely.
            # For PDF: fall back to regex on body lines.
            if suffix in (".docx", ".doc"):
                ref_to_claim = _build_docx_ref_claims(doc_path, ref_map)
            else:
                body_lines = [l.strip() for l in body_text.splitlines() if l.strip()]
                ref_to_claim = {
                    rn: _find_claim_for_ref(body_lines, rn)
                    for rn in ref_map
                }

            for ref_num, url in sorted(ref_map.items()):
                if url in seen:
                    continue
                seen.add(url)
                claim = ref_to_claim.get(ref_num, "")
                doi: Optional[str] = None
                if "doi.org/" in url:
                    doi = url.split("doi.org/", 1)[1]
                citations.append(Citation(
                    index=len(citations) + 1,
                    claim=claim,
                    raw_ref=url,
                    source_url=url,
                    doi=doi,
                    claim_source="body" if claim else "ref_only",
                ))

            logger.info("[citation_verifier] Extracted %d citations (numbered ref list) from %s",
                        len(citations), Path(doc_path).name)
            return citations

    # ── Path B: inline URL/DOI scan (fallback) ────────────────────────────────
    full_text = body_text + "\n" + ref_text
    sentences = _split_sentences(full_text)

    for idx, sentence in enumerate(sentences):
        for url in _URL_RE.findall(sentence):
            url = url.rstrip(".,;)")
            if url in seen:
                continue
            seen.add(url)
            citations.append(Citation(
                index=len(citations) + 1,
                claim=_build_claim(sentences, idx),
                raw_ref=url,
                source_url=url,
            ))

        for doi in _DOI_RE.findall(sentence):
            doi = doi.rstrip(".,;)")
            if doi in seen:
                continue
            seen.add(doi)
            citations.append(Citation(
                index=len(citations) + 1,
                claim=_build_claim(sentences, idx),
                raw_ref=doi,
                source_url=f"https://doi.org/{doi}",
                doi=doi,
            ))

    logger.info("[citation_verifier] Extracted %d citations (inline scan) from %s",
                len(citations), Path(doc_path).name)
    return citations


# ── Step 2: Source Resolution ─────────────────────────────────────────────────

# _resolve_doi, _core_fetch_by_doi, _extract_doi_from_page, _fetch_pdf_text,
# and _fetch_via_archives are imported from open_access_resolver at the top of this file.


async def _fetch_source_text(url: str) -> tuple[Optional[str], Optional[str]]:
    """
    Retrieve full text from a URL.
    Returns (text, page_title).

    Priority: trafilatura (article extraction) → existing BrowserTool.
    """
    # trafilatura is best for article/journal/news pages
    try:
        import trafilatura
        downloaded = await asyncio.to_thread(trafilatura.fetch_url, url)
        if downloaded:
            text = await asyncio.to_thread(
                trafilatura.extract,
                downloaded,
                include_comments=False,
                include_tables=True,
                no_fallback=False,
            )
            if text and len(text) > 200:
                meta = trafilatura.extract_metadata(downloaded)
                title = meta.title if meta else None
                logger.debug("[source] trafilatura: %d chars from %s", len(text), url)
                return text[:_MAX_SOURCE_CHARS], title
    except Exception as exc:
        logger.debug("[source] trafilatura failed for %s: %s", url, exc)

    # BrowserTool fallback (handles JS-rendered pages)
    try:
        from app.tools.browser import get_browser_tool
        result = await get_browser_tool().fetch_page(url)
        if "error" not in result:
            return result.get("text_content"), result.get("title")
    except Exception as exc:
        logger.debug("[source] BrowserTool failed for %s: %s", url, exc)

    return None, None




async def resolve_source(citation: Citation) -> Citation:
    """Resolve DOI → URL and fetch source full text. Sets resolved_url and source_text."""
    url = citation.source_url

    if citation.doi:
        resolved = await _resolve_doi(citation.doi)
        citation.resolved_url = resolved or f"https://doi.org/{citation.doi}"
        url = citation.resolved_url
    else:
        citation.resolved_url = url

    # Direct PDF links: download and extract text instead of browser fetch
    if url.lower().split("?")[0].endswith(".pdf"):
        text, title = await _fetch_pdf_text(url)
        citation.page_title = title
        citation.source_text = text
        if not text:
            citation.status = "unreachable"
            citation.error = f"Could not extract text from PDF: {url}"
            logger.warning("[source] PDF unreachable/unreadable: %s", url)
        return citation

    text, title = await _fetch_source_text(url)
    citation.page_title = title

    # Detect paywall / login / insufficient content pages — try archives before giving up
    blocked = False
    if not text:
        blocked = True
    elif len(text) < 400:
        blocked = True
    else:
        text_lower = text.lower()
        blocked = (
            "subscribe to" in text_lower
            or "sign in to read" in text_lower
            or "create a free account" in text_lower
            or "this content is available to subscribers" in text_lower
            or ("access denied" in text_lower and len(text) < 2000)
        )

    if blocked:
        logger.info("[source] Primary fetch blocked for %s — trying OA aggregators + archives…", url)

        # If we don't have a DOI yet, try to extract one from the page's HTML meta tags.
        # This lets us unlock the full OA resolution chain (CORE, Semantic Scholar, Unpaywall)
        # for plain-URL citations that the author didn't cite by DOI.
        doi_for_lookup = citation.doi
        if not doi_for_lookup:
            extracted_doi = await _extract_doi_from_page(url)
            if extracted_doi:
                doi_for_lookup = extracted_doi
                citation.doi = extracted_doi   # persist so it appears in the report
                logger.info("[source] DOI extracted from page meta: %s", extracted_doi)

        archive_text, archive_title = await _fetch_via_archives(url, doi=doi_for_lookup)
        if archive_text:
            citation.source_text = archive_text
            citation.page_title = archive_title or citation.page_title
            source_label = "CORE repository" if doi_for_lookup else "web archive"
            citation.error = f"Retrieved via {source_label} (primary source blocked/paywalled)"
            logger.info("[source] OA/archive fallback succeeded for %s", url)
            return citation

        # All OA sources and archives exhausted
        reason = "Paywall or insufficient content" if text else "Could not retrieve content"
        citation.status = "unreachable"
        citation.error = (
            f"{reason} — {url}. Tried: CORE API, Semantic Scholar, "
            "Wayback Machine, archive.ph — all failed."
        )
        logger.warning("[source] Unreachable after all OA + archive fallbacks: %s", url)
        return citation

    citation.source_text = text
    return citation


# ── Step 3: Quote Verification ────────────────────────────────────────────────

def _fuzzy_window_match(claim: str, source: str, window: int = 600) -> tuple[float, str]:
    """
    Find the best matching passage in `source` for `claim`.

    Uses rapidfuzz.fuzz.partial_ratio when available — this scores the shorter
    string against the best-aligned substring of the longer string, which is
    exactly right for matching a sentence-length claim against a full article.
    Falls back to a sliding difflib window when rapidfuzz is unavailable.
    """
    claim_lower = claim.lower().strip()
    if not claim_lower:
        return 0.0, ""

    try:
        from rapidfuzz import fuzz as _fuzz, process as _proc

        # partial_ratio: best alignment of claim inside any substring of source
        score = _fuzz.partial_ratio(claim_lower, source.lower()) / 100.0

        # Find the best matching window for display
        best_passage = ""
        best_chunk_score = 0.0
        step = window // 2
        for i in range(0, len(source), step):
            chunk = source[i: i + window]
            s = _fuzz.partial_ratio(claim_lower, chunk.lower()) / 100.0
            if s > best_chunk_score:
                best_chunk_score = s
                best_passage = chunk.strip()

        return score, best_passage

    except ImportError:
        import difflib

        best_score = 0.0
        best_passage = ""
        step = window // 2
        for i in range(0, len(source), step):
            chunk = source[i: i + window]
            s = difflib.SequenceMatcher(None, claim_lower, chunk.lower()).ratio()
            if s > best_score:
                best_score = s
                best_passage = chunk.strip()

        return best_score, best_passage


def _score_to_status(score: float) -> str:
    if score >= _SCORE_CONFIRMED:
        return "confirmed"
    if score >= _SCORE_PARTIAL:
        return "partial"
    if score >= _SCORE_UNCERTAIN:
        return "uncertain"
    return "hallucinated"


async def verify_quote(citation: Citation) -> Citation:
    """
    Fuzzy-match the claim against fetched source text, then run relational
    entity back-check for all scores below _SCORE_CONFIRMED.

    HALLUCINATED is only final when BOTH fuzzy AND relational find nothing —
    a score of 0.0 is expected for fact-integrated prose (not direct quotes).
    """
    if citation.status == "unreachable" or not citation.source_text:
        return citation

    score, passage = await asyncio.to_thread(
        _fuzzy_window_match, citation.claim, citation.source_text
    )
    citation.match_score = round(score, 3)
    citation.matched_passage = passage
    citation.status = _score_to_status(score)

    if score < _SCORE_CONFIRMED:
        verdict, note = await _verify_with_llm(citation)
        citation.llm_verdict = verdict or None
        citation.llm_note    = note or None
        if verdict == "CONFIRMED":
            citation.status = "confirmed"
        elif verdict == "PARTIAL":
            citation.status = "partial"
        elif verdict == "NOT_SUPPORTED":
            citation.status = "hallucinated" if score < _SCORE_UNCERTAIN else "uncertain"

    logger.info("[verify] Citation #%d  score=%.2f  status=%s  llm=%s",
                citation.index, score, citation.status, citation.llm_verdict or "—")
    return citation


async def verify_section(
    section_text: str,
    area_id: str,
    emit_fn=None,
) -> VerificationReport:
    """
    Verify citations in an in-memory section string (team pipeline gate mode).

    Skips file parsing, screenshots, and PDF report generation.
    LightRAG is the primary source — sprint agents already absorbed source content
    during their research phase. Falls back to live HTTP fetch if source not in graph.

    Returns a VerificationReport with citations and their statuses.
    """
    report = VerificationReport(doc_path=f"team:{area_id}", doc_name=area_id)

    sentences = _split_sentences(section_text)
    citations: list[Citation] = []
    seen: set[str] = set()

    for idx, sentence in enumerate(sentences):
        for url in _URL_RE.findall(sentence):
            url = url.rstrip(".,;)")
            if url not in seen:
                seen.add(url)
                citations.append(Citation(
                    index=len(citations) + 1,
                    claim=_build_claim(sentences, idx),
                    raw_ref=url,
                    source_url=url,
                ))
        for doi in _DOI_RE.findall(sentence):
            doi = doi.rstrip(".,;)")
            if doi not in seen:
                seen.add(doi)
                citations.append(Citation(
                    index=len(citations) + 1,
                    claim=_build_claim(sentences, idx),
                    raw_ref=doi,
                    source_url=f"https://doi.org/{doi}",
                    doi=doi,
                ))

    report.citations = citations

    for citation in citations:
        # Team mode: LightRAG-first source retrieval (warm from sprint research)
        source_id = citation.doi or citation.source_url or ""
        if source_id:
            try:
                from app.service.lightrag_service import LightRAGService
                lg = LightRAGService.get_instance()
                if lg._available:
                    res = await lg.query(
                        f"full text content from source {source_id}", mode="naive"
                    )
                    if res.get("success") and len(res.get("result", "")) > 300:
                        citation.source_text = res["result"]
                        citation.resolved_url = citation.source_url
            except Exception:
                pass

        # Fall back to live fetch if not yet in graph
        if not citation.source_text:
            citation = await resolve_source(citation)

        if citation.status != "unreachable":
            citation = await verify_quote(citation)

    logger.info("[verify_section] %s: %d citations checked", area_id, len(citations))
    return report


def _build_canvas_summary(report: VerificationReport) -> str:
    """Build a compact markdown summary of verification results for render_canvas."""
    s = report.summary
    lines = [
        f"**{s['total']} citations verified** for `{report.doc_name}`\n",
        f"- ✅ Confirmed: {s['confirmed']}",
        f"- 🟡 Partial: {s.get('partial', 0)}",
        f"- 🟠 Uncertain: {s.get('uncertain', 0)}",
        f"- ❌ Not found: {s.get('hallucinated', 0)}",
        f"- ⚫ Unreachable: {s.get('unreachable', 0)}",
    ]
    if report.doc_summary:
        lines += ["", f"**Document:** {report.doc_summary}"]

    flagged = [c for c in report.citations if c.status in ("hallucinated", "uncertain")]
    if flagged:
        lines += ["", "**Flagged citations:**"]
        for c in flagged[:10]:
            lines.append(
                f"- #{c.index} [{c.status.upper()}] score={c.match_score:.2f} — "
                f"{(c.claim or '')[:120]}…"
            )
        if len(flagged) > 10:
            lines.append(f"- …and {len(flagged) - 10} more")

    return "\n".join(lines)


# ── Step 4: Screenshot ────────────────────────────────────────────────────────

# Injected into the page to highlight the matched passage, scroll to it,
# and return the bounding box for cropped screenshot capture.
_HIGHLIGHT_JS = """
(passage) => {
    const needle = (passage || "").slice(0, 120).trim().toLowerCase();
    if (!needle) return { found: false };

    const walker = document.createTreeWalker(
        document.body, NodeFilter.SHOW_TEXT, null, false
    );
    let node;
    while ((node = walker.nextNode())) {
        if (node.textContent.toLowerCase().includes(needle)) {
            try {
                const parent = node.parentElement;
                if (!parent) continue;
                const mark = document.createElement('mark');
                mark.id = '__aura_cite_mark__';
                mark.style.cssText =
                    'background:#FFD700;padding:3px 5px;border-radius:3px;';
                parent.insertBefore(mark, node);
                mark.appendChild(node);
                mark.scrollIntoView({ behavior: 'instant', block: 'center' });
                const r = mark.getBoundingClientRect();
                return {
                    found: true,
                    x: Math.max(0, r.x - 120),
                    y: Math.max(0, r.y - 100),
                    width:  Math.min(1280, r.width  + 240),
                    height: Math.min(900,  r.height + 200)
                };
            } catch (_) {
                continue;
            }
        }
    }
    return { found: false };
}
"""


def _render_pdf_card(citation: Citation) -> bytes:
    """
    For PDF-sourced citations, render a clean dark card showing the matched
    passage text — used in place of a browser screenshot.
    """
    W, H = 1280, 460
    img = Image.new("RGB", (W, H), (15, 23, 42))
    draw = ImageDraw.Draw(img)

    color = _STATUS_RGB.get(citation.status, (148, 163, 184))
    draw.rectangle([(0, 0), (W, 44)], fill=color)
    draw.text((12, 12),
              f"  PDF Source  #{citation.index}  {citation.status.upper()}  "
              f"score: {citation.match_score:.2f}",
              fill=(255, 255, 255), font=_try_font(18))

    label_font = _try_font(11)
    body_font  = _try_font(14)

    draw.text((20, 58), "MATCHED PASSAGE:", fill=(100, 116, 139), font=label_font)
    passage = citation.matched_passage or "(no passage matched)"
    # Naive word-wrap at ~100 chars per line
    words, line, lines = passage.split(), "", []
    for w in words:
        if len(line) + len(w) + 1 > 100:
            lines.append(line)
            line = w
        else:
            line = (line + " " + w).strip()
    if line:
        lines.append(line)
    y = 78
    for ln in lines[:12]:
        draw.text((20, y), ln, fill=(203, 213, 225), font=body_font)
        y += 22

    draw.rectangle([(0, H - 30), (W, H)], fill=(30, 41, 59))
    url_text = (citation.resolved_url or citation.source_url or "")[:130]
    draw.text((12, H - 21), url_text, fill=(148, 163, 184), font=_try_font(12))

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


async def capture_screenshot(citation: Citation) -> Citation:
    """
    Navigate to the resolved source URL, highlight the matched passage,
    and capture a screenshot cropped to the highlighted quote region.
    Stores PNG bytes in citation.screenshot_png.
    """
    url = citation.resolved_url or citation.source_url
    if not url:
        return citation

    # PDF sources: render a text card instead of browser navigation
    if url.lower().split("?")[0].endswith(".pdf"):
        citation.screenshot_png = _render_pdf_card(citation)
        citation.highlight_found = bool(citation.matched_passage)
        return citation

    try:
        from playwright.async_api import async_playwright

        async with async_playwright() as pw:
            browser = await pw.chromium.launch(headless=True)
            ctx = await browser.new_context(
                viewport={"width": 1280, "height": 900},
                user_agent=(
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/120.0.0.0 Safari/537.36"
                ),
            )
            page = await ctx.new_page()
            try:
                await page.goto(url, wait_until="domcontentloaded", timeout=_BROWSER_TIMEOUT)
                # Brief pause for dynamic content
                await page.wait_for_timeout(1200)

                if citation.matched_passage and citation.status not in ("unreachable", "hallucinated"):
                    try:
                        result = await page.evaluate(_HIGHLIGHT_JS, citation.matched_passage)
                        await page.wait_for_timeout(400)

                        if isinstance(result, dict) and result.get("found"):
                            citation.highlight_found = True
                            clip = {
                                "x":      result["x"],
                                "y":      result["y"],
                                "width":  max(result["width"],  500),
                                "height": max(result["height"], 140),
                            }
                            citation.screenshot_png = await page.screenshot(
                                type="png", clip=clip
                            )
                        else:
                            citation.screenshot_png = await page.screenshot(type="png")
                    except Exception:
                        citation.screenshot_png = await page.screenshot(type="png")
                else:
                    citation.screenshot_png = await page.screenshot(type="png")

                logger.info("[screenshot] Citation #%d captured (highlight=%s) (%s)",
                            citation.index, citation.highlight_found, url[:80])
            finally:
                await page.close()
                await ctx.close()
                await browser.close()

    except Exception as exc:
        logger.warning("[screenshot] Citation #%d failed: %s", citation.index, exc)
        citation.error = (citation.error or "") + f" | Screenshot error: {exc}"

    return citation


# ── Step 4b: Screenshot Annotation ───────────────────────────────────────────

_STATUS_RGB = {
    "confirmed":    (34,  197,  94),   # green
    "partial":      (251, 191,  36),   # amber
    "uncertain":    (249, 115,  22),   # orange
    "hallucinated": (239,  68,  68),   # red
    "unreachable":  (148, 163, 184),   # slate
    "pending":      (148, 163, 184),
}


def _try_font(size: int) -> ImageFont.ImageFont:
    for name in ("arial.ttf", "Arial.ttf", "DejaVuSans.ttf"):
        try:
            return ImageFont.truetype(name, size)
        except Exception:
            pass
    return ImageFont.load_default()


def annotate_screenshot(citation: Citation) -> Citation:
    """
    Overlay a status badge (top bar) and source URL (bottom bar) onto the
    raw screenshot. Stores result in citation.annotated_png.
    """
    if not citation.screenshot_png:
        return citation

    try:
        img = Image.open(io.BytesIO(citation.screenshot_png)).convert("RGB")
        draw = ImageDraw.Draw(img)
        color = _STATUS_RGB.get(citation.status, (148, 163, 184))

        # Top bar
        locate_note = "" if citation.highlight_found else "  ⚠ quote not located on page"
        badge_text = (
            f"  #{citation.index}  {citation.status.upper()}"
            f"  score: {citation.match_score:.2f}"
            f"{locate_note}"
        )
        draw.rectangle([(0, 0), (img.width, 38)], fill=color)
        draw.text((10, 9), badge_text, fill=(255, 255, 255), font=_try_font(18))

        # Bottom bar
        url_text = (citation.resolved_url or citation.source_url or "")[:130]
        bar_y = img.height - 30
        draw.rectangle([(0, bar_y), (img.width, img.height)], fill=(30, 41, 59))
        draw.text((10, bar_y + 7), url_text, fill=(203, 213, 225), font=_try_font(13))

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        citation.annotated_png = buf.getvalue()

    except Exception as exc:
        logger.warning("[annotate] Citation #%d failed: %s", citation.index, exc)
        citation.annotated_png = citation.screenshot_png

    return citation


async def _summarize_document(body_text: str, doc_name: str) -> str:
    """
    Generate a short document overview using the local Ollama workhorse.
    Returns a 2–3 sentence summary for the PDF cover page, or "" on failure.
    """
    from app.service.ollama_service import get_ollama_service
    svc = get_ollama_service()
    if svc is None:
        return ""
    try:
        return await svc.chat(
            messages=[{"role": "user", "content": (
                f"Summarize this excerpt from '{doc_name}' in 2–3 sentences. "
                "What is the document about and what is its main argument?\n\n"
                f"\"\"\"\n{body_text[:3000].strip()}\n\"\"\""
            )}],
            temperature=0.3, max_tokens=150,
        )
    except Exception as exc:
        logger.debug("[doc_summary] Workhorse failed: %s", exc)
        return ""


async def _verify_with_llm(citation: "Citation") -> tuple[str, str]:
    """
    Entity-level semantic verification via local workhorse (DeepSeek-R1 14B).

    Called for ALL fuzzy scores below _SCORE_CONFIRMED (0.75).
    A score of 0.0 is expected for fact-integrated prose — HALLUCINATED is only
    final when BOTH fuzzy AND relational find nothing.

    Vacuum constraint: only sees the claim, the source excerpt, and optionally
    scoped prior content from this specific source. No AURA general memory injected.
    """
    from app.service.ollama_service import get_ollama_service
    svc = get_ollama_service()
    if svc is None:
        return "", ""

    graph_context = ""
    source_id = citation.doi or citation.resolved_url or citation.source_url or ""
    if source_id:
        try:
            from app.service.lightrag_service import LightRAGService
            lg = LightRAGService.get_instance()
            if lg._available:
                res = await lg.query(f"content from source {source_id}", mode="local")
                if res.get("success") and len(res.get("result", "")) > 200:
                    graph_context = "\nPrior verified content from this source:\n" + res["result"][:600]
        except Exception as exc:
            logger.debug("[llm_verify] Graph context failed: %s", exc)

    excerpt = (citation.source_text or "")[:2000].strip()
    prompt = (
        "Check whether this cited source supports the claim.\n"
        "Focus on facts, entities, statistics — not exact wording.\n\n"
        f"CLAIM:\n\"\"\"\n{citation.claim}\n\"\"\"\n\n"
        f"SOURCE TEXT:\n\"\"\"\n{excerpt}\n\"\"\""
        f"{graph_context}\n\n"
        "Base your answer only on the source content above.\n"
        'Respond with JSON: {"verdict": "CONFIRMED" | "PARTIAL" | "NOT_SUPPORTED", '
        '"note": "<one sentence on entity match or mismatch>"}'
    )
    try:
        result = await svc.chat_json(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
        )
        v, n = result.get("verdict", ""), result.get("note", "")
        if v in ("CONFIRMED", "PARTIAL", "NOT_SUPPORTED"):
            return v, n
    except Exception as exc:
        logger.debug("[llm_verify] Workhorse failed: %s", exc)
    return "", ""


def _build_score_explanation(c: "Citation") -> str:
    """
    Return a plain-English explanation of how and why this citation received its score.
    Combines fuzzy match logic, LLM entity verdict, source access notes, and any warnings.
    """
    parts: list[str] = []

    # String similarity
    pct = f"{c.match_score:.0%}"
    if c.match_score >= _SCORE_CONFIRMED:
        parts.append(
            f"String similarity {pct}: High lexical overlap — the claim's key terms "
            "appear verbatim or near-verbatim in the retrieved source text."
        )
    elif c.match_score >= _SCORE_PARTIAL:
        parts.append(
            f"String similarity {pct}: Moderate overlap — key terms partially match, "
            "suggesting the claim may be paraphrased or aggregated from multiple source sentences."
        )
    elif c.match_score >= _SCORE_UNCERTAIN:
        parts.append(
            f"String similarity {pct}: Low overlap — claim language differs substantially "
            "from source text; the statistic or assertion may be drawn from a section "
            "not captured in the retrieved excerpt."
        )
    elif c.status != "unreachable":
        parts.append(
            f"String similarity {pct}: Minimal overlap — the claim's specific language "
            "was not found in the retrieved source content."
        )

    # LLM relational entity verdict
    if c.llm_verdict and c.llm_verdict != "PARTIAL":
        verdict_label = {
            "CONFIRMED": "Semantically confirmed",
            "NOT_SUPPORTED": "Not semantically supported",
        }.get(c.llm_verdict, c.llm_verdict)
        note = f" ({c.llm_note})" if c.llm_note else ""
        parts.append(f"LLM entity analysis: {verdict_label}{note}.")
    elif c.llm_verdict == "PARTIAL" and c.llm_note:
        parts.append(f"LLM analysis: {c.llm_note}.")

    # Source access conditions
    if c.status == "unreachable":
        parts.append(
            "Source unreachable: Content could not be retrieved — the page may be "
            "paywalled, behind a login, or no longer available. Archive fallback also failed."
        )
    elif c.error and ("archive" in c.error.lower() or "paywall" in c.error.lower()):
        parts.append(
            "Source access limited: Primary URL was paywalled or blocked. "
            "Content retrieved via public web archive — coverage may be incomplete."
        )

    # Extraction quality warnings
    if c.claim_source == "ref_only":
        parts.append(
            "Extraction note: No inline citation marker was found in the document body for "
            "this reference. The claim context shown is derived from the reference list entry "
            "only — manual review is recommended."
        )

    if not c.highlight_found and c.status not in ("unreachable", "hallucinated", "pending"):
        parts.append(
            "Display note: The matched passage could not be highlighted on the live source page — "
            "the text may be rendered differently in the browser than in the extracted content."
        )

    return " ".join(parts) if parts else "No scoring detail available."


# ── Step 5: PDF Report ────────────────────────────────────────────────────────

def _hex(status: str) -> str:
    r, g, b = _STATUS_RGB.get(status, (148, 163, 184))
    return f"#{r:02X}{g:02X}{b:02X}"


def build_pdf_report(report: VerificationReport, output_path: str) -> str:
    """
    Generate the final citation report PDF.

    Structure:
      • Cover page  — report header, document overview (LLM-generated), results summary,
                      verification methodology
      • Per-citation — status badge, metadata, claim, matched passage, scoring
                       logic (LLM entity check + fuzzy), screenshot cropped to the quote
    """
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable, Image as RLImage, PageBreak,
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )
    from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate

    PAGE_W_FULL = letter[0]
    PAGE_H_FULL = letter[1]
    MARGIN = 0.75 * inch
    PAGE_W = PAGE_W_FULL - 2 * MARGIN

    # ── Colour palette ─────────────────────────────────────────────────────────
    NAVY       = colors.HexColor("#0F2645")
    NAVY_MID   = colors.HexColor("#1E3A5F")
    SLATE      = colors.HexColor("#64748B")
    SLATE_LT   = colors.HexColor("#CBD5E1")
    BG_ROW     = colors.HexColor("#F8FAFC")
    BG_LABEL   = colors.HexColor("#F1F5F9")
    WHITE      = colors.white

    STATUS_COLOR = {
        "confirmed":    colors.HexColor("#16A34A"),
        "partial":      colors.HexColor("#D97706"),
        "uncertain":    colors.HexColor("#EA580C"),
        "hallucinated": colors.HexColor("#DC2626"),
        "unreachable":  colors.HexColor("#64748B"),
        "pending":      colors.HexColor("#94A3B8"),
    }

    # ── Page callback — running header / footer ────────────────────────────────
    _doc_label = Path(report.doc_path).name
    _gen_date  = report.generated_at[:10]

    def _on_page(canvas, doc):
        canvas.saveState()
        # Thin navy top rule
        canvas.setStrokeColor(NAVY_MID)
        canvas.setLineWidth(1)
        canvas.line(MARGIN, PAGE_H_FULL - 0.45 * inch,
                    PAGE_W_FULL - MARGIN, PAGE_H_FULL - 0.45 * inch)
        # Header text
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(SLATE)
        canvas.drawString(MARGIN, PAGE_H_FULL - 0.38 * inch, "CITATION VERIFICATION REPORT")
        canvas.drawRightString(PAGE_W_FULL - MARGIN, PAGE_H_FULL - 0.38 * inch,
                               _doc_label[:80])
        # Footer rule
        canvas.line(MARGIN, 0.55 * inch, PAGE_W_FULL - MARGIN, 0.55 * inch)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(SLATE)
        canvas.drawString(MARGIN, 0.38 * inch, f"Gleipnir Consulting  |  {_gen_date}")
        canvas.drawRightString(PAGE_W_FULL - MARGIN, 0.38 * inch,
                               f"Page {doc.page}")
        canvas.restoreState()

    # ── Styles ─────────────────────────────────────────────────────────────────
    base = getSampleStyleSheet()

    sty_title = ParagraphStyle(
        "RTitle", parent=base["Title"],
        fontSize=26, leading=30, textColor=WHITE,
        alignment=TA_LEFT, spaceAfter=0,
    )
    sty_subtitle = ParagraphStyle(
        "RSubtitle", parent=base["Normal"],
        fontSize=11, textColor=colors.HexColor("#93C5FD"),
        alignment=TA_LEFT, spaceAfter=0,
    )
    sty_h2 = ParagraphStyle(
        "RH2", parent=base["Heading2"],
        fontSize=12, textColor=NAVY_MID, spaceBefore=6, spaceAfter=3,
    )
    sty_h3 = ParagraphStyle(
        "RH3", parent=base["Heading3"],
        fontSize=10, textColor=SLATE, spaceBefore=4, spaceAfter=2,
    )
    sty_body = ParagraphStyle(
        "RBody", parent=base["Normal"],
        fontSize=10, leading=14, spaceAfter=4,
    )
    sty_label = ParagraphStyle(
        "RLabel", parent=base["Normal"],
        fontSize=8.5, textColor=SLATE, spaceAfter=1,
    )
    sty_mono = ParagraphStyle(
        "RMono", parent=base["Code"],
        fontSize=9, leading=13, spaceAfter=4,
        backColor=colors.HexColor("#F8FAFC"),
    )
    sty_warn = ParagraphStyle(
        "RWarn", parent=base["Normal"],
        fontSize=9, textColor=colors.HexColor("#C2410C"),
        spaceAfter=3,
    )
    sty_logic = ParagraphStyle(
        "RLogic", parent=base["Normal"],
        fontSize=9, leading=13, textColor=colors.HexColor("#374151"),
        backColor=colors.HexColor("#F0F9FF"),
        spaceAfter=4, leftIndent=6, rightIndent=6,
    )
    sty_llm_verdict = ParagraphStyle(
        "RLLMVerdict", parent=base["Normal"],
        fontSize=9, leading=13, spaceAfter=3,
    )
    sty_method = ParagraphStyle(
        "RMethod", parent=base["Normal"],
        fontSize=9, leading=14, textColor=colors.HexColor("#374151"),
        spaceAfter=3,
    )

    # ── Doc setup ──────────────────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=0.65 * inch, bottomMargin=0.75 * inch,
        onPage=_on_page,
    )

    summary = report.summary
    story   = []

    # ══ COVER PAGE ═════════════════════════════════════════════════════════════

    # Navy masthead block
    masthead_data = [[
        Paragraph("Citation Verification Report", sty_title),
        Paragraph(
            f"{_doc_label}<br/>"
            f"<font size='9' color='#93C5FD'>{_gen_date}</font>",
            sty_subtitle,
        ),
    ]]
    masthead_tbl = Table(masthead_data, colWidths=[PAGE_W])
    masthead_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, -1), NAVY),
        ("LEFTPADDING",  (0, 0), (-1, -1), 18),
        ("RIGHTPADDING", (0, 0), (-1, -1), 18),
        ("TOPPADDING",   (0, 0), (-1, -1), 20),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 20),
        ("ROUNDEDCORNERS", (0, 0), (-1, -1), 4),
    ]))
    story.append(masthead_tbl)
    story.append(Spacer(1, 0.25 * inch))

    # ── Document overview ──────────────────────────────────────────────────────
    story.append(Paragraph("Document Overview", sty_h2))
    story.append(HRFlowable(width="100%", thickness=1, color=SLATE_LT))
    story.append(Spacer(1, 0.08 * inch))

    if report.doc_summary:
        story.append(Paragraph(report.doc_summary, sty_body))
    else:
        story.append(Paragraph(
            f"Document: <b>{_doc_label}</b>", sty_body,
        ))
    story.append(Spacer(1, 0.2 * inch))

    # ── Results summary ────────────────────────────────────────────────────────
    story.append(Paragraph("Verification Results", sty_h2))
    story.append(HRFlowable(width="100%", thickness=1, color=SLATE_LT))
    story.append(Spacer(1, 0.08 * inch))

    STATUS_LABELS = [
        ("confirmed",    "Confirmed",    "≥ 75% similarity — source directly supports the claim."),
        ("partial",      "Partial",      "45–74% — claim is supported but paraphrased or summarised."),
        ("uncertain",    "Uncertain",    "20–44% — weak match; claim may be synthesised from multiple sources."),
        ("hallucinated", "Not Found",    "< 20% — claim language absent from retrieved source content."),
        ("unreachable",  "Unreachable",  "Source could not be retrieved (paywall, login, or unavailable)."),
    ]
    results_rows = [["Status", "Count", "Meaning"]]
    for key, label, meaning in STATUS_LABELS:
        results_rows.append([
            Paragraph(
                f'<font color="{_hex(key)}"><b>{label}</b></font>',
                ParagraphStyle("RS", parent=base["Normal"], fontSize=9),
            ),
            str(summary.get(key, 0)),
            Paragraph(meaning,
                      ParagraphStyle("RM", parent=base["Normal"], fontSize=9)),
        ])
    results_rows.append([
        Paragraph("<b>Total</b>",
                  ParagraphStyle("RT", parent=base["Normal"], fontSize=9)),
        str(summary["total"]),
        "",
    ])

    res_tbl = Table(results_rows, colWidths=[1.3 * inch, 0.55 * inch, PAGE_W - 1.85 * inch])
    res_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0),  NAVY),
        ("TEXTCOLOR",    (0, 0), (-1, 0),  WHITE),
        ("FONTSIZE",     (0, 0), (-1, -1), 9),
        ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2),
         [WHITE, BG_ROW]),
        ("BACKGROUND",   (0, -1), (-1, -1), BG_LABEL),
        ("FONTNAME",     (0, -1), (-1, -1), "Helvetica-Bold"),
        ("GRID",         (0, 0), (-1, -1), 0.5, SLATE_LT),
        ("LEFTPADDING",  (0, 0), (-1, -1), 8),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(res_tbl)
    story.append(Spacer(1, 0.25 * inch))

    # ── Verification methodology ───────────────────────────────────────────────
    story.append(Paragraph("Verification Methodology", sty_h2))
    story.append(HRFlowable(width="100%", thickness=1, color=SLATE_LT))
    story.append(Spacer(1, 0.08 * inch))

    method_lines = [
        "<b>Step 1 — Citation extraction:</b> The document is parsed to identify all numbered "
        "references. For DOCX files, python-docx superscript run metadata is used to map each "
        "reference number to the exact body paragraph that contains the inline citation marker, "
        "ensuring the claim reflects what the author was asserting at that point.",

        "<b>Step 2 — Source resolution:</b> Each URL is fetched using trafilatura article "
        "extraction. DOIs are resolved through CrossRef and Unpaywall to obtain open-access "
        "full-text where available. Paywalled sources trigger an automatic fallback to the "
        "Wayback Machine and archive.ph.",

        "<b>Step 3 — String similarity:</b> The extracted claim is fuzzy-matched against the "
        "retrieved source text using rapidfuzz partial_ratio, which finds the best alignment "
        "of the shorter claim within the longer source. A score of 1.0 = verbatim match; "
        "lower scores reflect paraphrase, summarisation, or different terminology.",

        "<b>Step 4 — Relational entity verification:</b> For all citations scoring below 0.75, "
        "the local workhorse LLM performs an entity-level back-check — extracting facts, "
        "statistics, and named entities from the claim and cross-referencing them against the "
        "source content. This catches fact-integrated prose where the author has absorbed a "
        "statistic into their own writing rather than quoting directly. HALLUCINATED is only "
        "assigned when both fuzzy matching and relational checking find no correspondence. "
        "LLM verdicts can upgrade or downgrade the string-similarity result.",

        "<b>Step 5 — Screenshot:</b> For web sources, a browser navigates to the source, "
        "injects a highlight on the matched passage, and captures a screenshot cropped to "
        "the highlighted region. For PDF sources, a text card is rendered showing the matched "
        "passage. Unreachable or paywalled sources receive a status note in place of a screenshot.",
    ]
    for line in method_lines:
        story.append(Paragraph(line, sty_method))
        story.append(Spacer(1, 0.04 * inch))

    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(
        "<i>Score thresholds: Confirmed ≥ 0.75 | Partial 0.45–0.74 | "
        "Uncertain 0.20–0.44 | Not Found &lt; 0.20. "
        "LLM relational entity check may adjust final status up or down for scores below 0.75.</i>",
        sty_label,
    ))
    story.append(PageBreak())

    # ══ CITATION SECTIONS ══════════════════════════════════════════════════════

    for c in report.citations:
        status_color = STATUS_COLOR.get(c.status, colors.HexColor("#94A3B8"))
        status_hex   = _hex(c.status)

        # ── Citation header banner ─────────────────────────────────────────────
        header_data = [[
            Paragraph(
                f'Citation #{c.index}',
                ParagraphStyle("CHNum", parent=base["Normal"],
                               fontSize=11, textColor=WHITE,
                               fontName="Helvetica-Bold"),
            ),
            Paragraph(
                f'<b>{c.status.upper()}</b>  '
                f'<font size="10">score: {c.match_score:.2f}</font>',
                ParagraphStyle("CHStatus", parent=base["Normal"],
                               fontSize=11, textColor=WHITE,
                               alignment=TA_CENTER),
            ),
            Paragraph(
                (c.page_title or "")[:60],
                ParagraphStyle("CHTitle", parent=base["Normal"],
                               fontSize=8.5, textColor=colors.HexColor("#BFDBFE")),
            ),
        ]]
        hdr_tbl = Table(
            header_data,
            colWidths=[1.4 * inch, 2.0 * inch, PAGE_W - 3.4 * inch],
        )
        hdr_tbl.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (-1, -1), status_color),
            ("LEFTPADDING",  (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING",   (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 8),
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(hdr_tbl)
        story.append(Spacer(1, 0.1 * inch))

        # ── Metadata ──────────────────────────────────────────────────────────
        url_display = (c.resolved_url or c.source_url or "—")[:120]
        meta_rows = []
        if c.page_title:
            meta_rows.append(["Source Title", c.page_title[:100]])
        meta_rows.append(["URL", url_display])
        if c.doi:
            meta_rows.append(["DOI", c.doi[:80]])
        meta_rows.append(["Match Score", f"{c.match_score:.3f}"])
        if c.llm_verdict:
            meta_rows.append(["LLM Verdict", c.llm_verdict])

        meta_tbl = Table(
            meta_rows,
            colWidths=[1.3 * inch, PAGE_W - 1.3 * inch],
        )
        meta_tbl.setStyle(TableStyle([
            ("BACKGROUND",   (0, 0), (0, -1), BG_LABEL),
            ("FONTSIZE",     (0, 0), (-1, -1), 9),
            ("FONTNAME",     (0, 0), (0, -1),  "Helvetica-Bold"),
            ("GRID",         (0, 0), (-1, -1), 0.4, SLATE_LT),
            ("ROWBACKGROUNDS", (0, 0), (-1, -1), [WHITE, BG_ROW]),
            ("LEFTPADDING",  (0, 0), (-1, -1), 6),
            ("TOPPADDING",   (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
            ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(meta_tbl)
        story.append(Spacer(1, 0.1 * inch))

        # ── Claim ─────────────────────────────────────────────────────────────
        if c.claim_source == "ref_only":
            story.append(Paragraph(
                "⚠ No inline citation marker found in document body — "
                "claim context derived from reference list only. Manual review required.",
                sty_warn,
            ))
        story.append(Paragraph("CLAIM FROM DOCUMENT", sty_label))
        story.append(Paragraph(
            f'"{(c.claim or "—")[:700]}"', sty_mono,
        ))
        story.append(Spacer(1, 0.08 * inch))

        # ── Best matching passage ─────────────────────────────────────────────
        if c.matched_passage:
            story.append(Paragraph("BEST MATCHING PASSAGE IN SOURCE", sty_label))
            story.append(Paragraph(c.matched_passage[:500], sty_mono))
            story.append(Spacer(1, 0.08 * inch))

        # ── Scoring logic ─────────────────────────────────────────────────────
        story.append(Paragraph("SCORING LOGIC", sty_label))
        story.append(Paragraph(_build_score_explanation(c), sty_logic))
        story.append(Spacer(1, 0.06 * inch))

        # ── LLM relational entity check ────────────────────────────────────────
        if c.llm_verdict:
            llm_color = {
                "CONFIRMED":     "#15803D",
                "PARTIAL":       "#B45309",
                "NOT_SUPPORTED": "#B91C1C",
            }.get(c.llm_verdict, "#64748B")
            story.append(Paragraph(
                f'<font color="{llm_color}"><b>● LLM Verdict: {c.llm_verdict}</b></font>'
                + (f"  —  {c.llm_note}" if c.llm_note else ""),
                sty_llm_verdict,
            ))
            story.append(Spacer(1, 0.06 * inch))

        # ── Access / error note ───────────────────────────────────────────────
        if c.error:
            story.append(Paragraph(f"Access note: {c.error}", sty_label))
            story.append(Spacer(1, 0.04 * inch))

        # ── Screenshot ────────────────────────────────────────────────────────
        img_bytes = c.annotated_png or c.screenshot_png
        if img_bytes:
            try:
                buf = io.BytesIO(img_bytes)
                pil = Image.open(buf)
                w, h = pil.size
                aspect = h / w
                disp_w = min(PAGE_W, 6.5 * inch)
                disp_h = disp_w * aspect
                if disp_h > 4.5 * inch:
                    disp_h = 4.5 * inch
                    disp_w = disp_h / aspect
                buf.seek(0)
                story.append(Paragraph("SOURCE SCREENSHOT — MATCHED PASSAGE", sty_label))
                story.append(Spacer(1, 0.04 * inch))
                story.append(RLImage(buf, width=disp_w, height=disp_h))
            except Exception as exc:
                logger.warning("[pdf] Image embed failed for #%d: %s", c.index, exc)

        story.append(Spacer(1, 0.2 * inch))
        story.append(PageBreak())

    doc.build(story)
    logger.info("[pdf_report] Written: %s", output_path)
    return output_path


# ── Orchestrator ──────────────────────────────────────────────────────────────

async def run_verification(
    doc_path: str,
    output_dir: Optional[str] = None,
    emit_fn: Optional[Callable] = None,
) -> VerificationReport:
    """
    Run the full citation verification pipeline on a PDF or DOCX file.

    Parameters
    ----------
    doc_path : str
        Absolute path to the input document.
    output_dir : str, optional
        Directory for the PDF report. Defaults to ~/.aura/citations/.
    emit_fn : async callable(event, payload), optional
        SSE progress callback compatible with AURA's interface agent.

    Returns
    -------
    VerificationReport
        Contains all Citation objects and the path to the generated PDF.
    """
    doc_path = str(Path(doc_path).resolve())

    if output_dir is None:
        import os as _os
        output_dir = (
            _os.environ.get("CITATION_OUTPUT_DIR")
            or str(Path.home() / ".aura" / "citations")
        )
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    report = VerificationReport(
        doc_path=doc_path,
        doc_name=Path(doc_path).stem,
    )

    async def _emit(event: str, payload: dict):
        if emit_fn:
            try:
                await emit_fn(event, payload)
            except Exception:
                pass

    # Step 1 — Extract
    await _emit("citation_progress", {"step": 1, "message": "Extracting citations…"})
    try:
        citations = await asyncio.to_thread(extract_citations, doc_path)
    except Exception as exc:
        logger.error("[verification] Extraction failed: %s", exc)
        await _emit("citation_error", {"step": 1, "error": str(exc)})
        return report

    if not citations:
        await _emit("citation_progress", {
            "step": 1, "done": True,
            "message": "No citations found in document.",
        })
        return report

    report.citations = citations

    # Generate document summary via local workhorse (non-blocking — runs before source fetches)
    try:
        suffix = Path(doc_path).suffix.lower()
        if suffix == ".pdf":
            body_text, _ = _split_body_refs_pdf(doc_path)
        else:
            body_text, _ = _split_body_refs_docx(doc_path)
        report.doc_summary = await _summarize_document(body_text, report.doc_name)
    except Exception as exc:
        logger.debug("[verification] Doc summary failed (non-fatal): %s", exc)
    await _emit("citation_progress", {
        "step": 1, "done": True,
        "message": f"Found {len(citations)} citation(s).",
    })

    # Steps 2–4 — per citation
    for citation in report.citations:
        await _emit("citation_progress", {
            "step": 2, "citation": citation.index,
            "message": f"Resolving citation #{citation.index}…",
        })

        citation = await resolve_source(citation)

        if citation.status != "unreachable":
            await _emit("citation_progress", {
                "step": 3, "citation": citation.index,
                "message": f"Verifying quote #{citation.index}…",
            })
            citation = await verify_quote(citation)

        await _emit("citation_progress", {
            "step": 4, "citation": citation.index,
            "message": f"Capturing screenshot #{citation.index}…",
        })
        citation = await capture_screenshot(citation)
        citation = annotate_screenshot(citation)

        await _emit("citation_complete", {
            "citation": citation.index,
            "status": citation.status,
            "score": citation.match_score,
        })

    # Step 5 — Build PDF
    await _emit("citation_progress", {"step": 5, "message": "Building PDF report…"})
    pdf_path = str(
        Path(output_dir) / f"{report.doc_name}_citation_report.pdf"
    )
    try:
        await asyncio.to_thread(build_pdf_report, report, pdf_path)
        report.pdf_path = pdf_path
        await _emit("citation_complete", {
            "step": 5, "done": True,
            "pdf_path": pdf_path,
            "summary": report.summary,
        })
    except Exception as exc:
        logger.error("[verification] PDF build failed: %s", exc)
        await _emit("citation_error", {"step": 5, "error": str(exc)})

    return report


# ── Batch: local folder ───────────────────────────────────────────────────────

_SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


async def run_folder_verification(
    folder_path: str,
    output_dir: Optional[str] = None,
    emit_fn: Optional[Callable] = None,
) -> list[VerificationReport]:
    """
    Run citation verification on every PDF/DOCX in a local folder.

    Documents are processed sequentially (screenshots are the bottleneck;
    concurrency here adds browser contention without meaningful speedup).

    Parameters
    ----------
    folder_path : str
        Absolute path to the folder containing documents.
    output_dir : str, optional
        Where to write per-document PDF reports.
        Defaults to ~/.aura/citations/<folder_name>/.
    emit_fn : async callable(event, payload), optional
        SSE progress callback.

    Returns
    -------
    list[VerificationReport]
        One report per document processed.
    """
    folder = Path(folder_path).resolve()
    if not folder.is_dir():
        raise ValueError(f"Not a directory: {folder}")

    docs = sorted(
        p for p in folder.iterdir()
        if p.is_file() and p.suffix.lower() in _SUPPORTED_EXTENSIONS
    )

    if not docs:
        logger.warning("[folder_batch] No PDF/DOCX files found in %s", folder)
        if emit_fn:
            await emit_fn("citation_batch_update", {
                "folder": str(folder),
                "total": 0,
                "message": "No PDF or DOCX documents found in the specified folder.",
            })
        return []

    if output_dir is None:
        output_dir = str(Path.home() / ".aura" / "citations" / folder.name)
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    if emit_fn:
        await emit_fn("citation_batch_start", {
            "folder": str(folder),
            "total": len(docs),
            "documents": [d.name for d in docs],
        })

    logger.info("[folder_batch] Processing %d documents in %s", len(docs), folder)
    reports: list[VerificationReport] = []

    for idx, doc_path in enumerate(docs, start=1):
        if emit_fn:
            await emit_fn("citation_batch_update", {
                "current": idx,
                "total": len(docs),
                "document": doc_path.name,
                "message": f"Processing {doc_path.name} ({idx}/{len(docs)})…",
            })

        try:
            report = await run_verification(
                doc_path=str(doc_path),
                output_dir=output_dir,
                emit_fn=emit_fn,
            )
            reports.append(report)
        except Exception as exc:
            logger.error("[folder_batch] Failed on %s: %s", doc_path.name, exc)
            if emit_fn:
                await emit_fn("citation_batch_error", {
                    "document": doc_path.name,
                    "error": str(exc),
                })

    if emit_fn:
        totals = {
            "confirmed": 0, "partial": 0, "uncertain": 0,
            "hallucinated": 0, "unreachable": 0,
        }
        for r in reports:
            for k in totals:
                totals[k] += r.summary.get(k, 0)

        await emit_fn("citation_batch_complete", {
            "folder": str(folder),
            "documents_processed": len(reports),
            "output_dir": output_dir,
            "totals": totals,
        })

    logger.info("[folder_batch] Done. %d/%d documents completed.", len(reports), len(docs))
    return reports


# ─────────────────────────────────────────────────────────────────────────────
# MCP TOOL INTERFACE — auto-registered by _mcp_wrapper.load_all_tools()
# ─────────────────────────────────────────────────────────────────────────────

TOOL_DEF = {
    "name": "citation_verifier",
    "description": (
        "Verify citations in a policy PDF or DOCX. Extracts claims using DOCX superscript "
        "metadata or PDF text; resolves sources via Semantic Scholar, Unpaywall, CrossRef, "
        "and CORE (200M+ OA records); fuzzy-matches claims with rapidfuzz; runs entity-level "
        "relational back-check via local workhorse LLM for ambiguous scores (handles "
        "fact-integrated prose where authors absorb statistics rather than quoting directly); "
        "captures screenshots cropped to the matched quote; and produces an annotated PDF "
        "report with scoring logic. Paywalled sources automatically trigger the "
        "open_access_resolver fallback chain (CORE → Semantic Scholar OA PDF → Wayback "
        "Machine → archive.ph). Returns per-citation status: confirmed, partial, uncertain, "
        "hallucinated, or unreachable."
    ),
    "inputSchema": {
        "type": "object",
        "properties": {
            "document_path": {
                "type": "string",
                "description": "Absolute path to the PDF or DOCX file to verify",
            },
            "output_dir": {
                "type": "string",
                "description": (
                    "Directory for the PDF report. "
                    "Defaults to CITATION_OUTPUT_DIR env var or ~/.aura/citations/"
                ),
            },
        },
        "required": ["document_path"],
    },
}


async def tool_handler(inputs: dict) -> dict:
    """MCP-compatible wrapper around run_verification() (external document mode)."""
    import os as _os
    doc_path = inputs.get("document_path", "")
    if not doc_path:
        return {"error": "document_path is required"}

    output_dir = inputs.get("output_dir") or _os.environ.get("CITATION_OUTPUT_DIR")

    # Wire SSE emit through AURA's chat controller
    emit_fn = None
    try:
        from app.controller.chat_controller import _emit as _aura_emit
        async def emit_fn(event: str, payload: dict):
            await _aura_emit(event, payload)
    except ImportError:
        pass

    if emit_fn:
        await emit_fn("agent_update", {
            "agent_id": "citation_verifier",
            "status":   "working",
            "summary":  f"Verifying citations: {Path(doc_path).name}",
        })

    try:
        report = await run_verification(doc_path, output_dir=output_dir, emit_fn=emit_fn)

        # Persist results to SQLite (non-fatal)
        run_id = None
        try:
            from app.service.citation_results_service import get_citation_results_service
            run_id = await asyncio.to_thread(
                get_citation_results_service().save_result, report
            )
            logger.info("[citation_verifier] Run persisted: %s", run_id)
        except Exception as db_exc:
            logger.warning("[citation_verifier] DB persist failed (non-fatal): %s", db_exc)

        # Absorb verified sources into LightRAG (warms future verification runs)
        for c in report.citations:
            if c.source_text and len(c.source_text) > 200:
                sid = c.doi or c.resolved_url or c.source_url or ""
                if sid:
                    try:
                        from app.service.lightrag_service import LightRAGService
                        LightRAGService.get_instance().enqueue_ingest(
                            c.source_text, f"source:{sid}", "citation_source"
                        )
                    except Exception:
                        pass

        # Absorb the full verification record into LightRAG
        try:
            from app.service.lightrag_service import LightRAGService
            lines = [f"# Citation Verification: {report.doc_name}", ""]
            if report.doc_summary:
                lines += [report.doc_summary, ""]
            for c in report.citations:
                lines.append(
                    f"[{c.status.upper()}] score={c.match_score:.2f} source={c.source_url}"
                )
                lines.append(f"Claim: {c.claim}")
                if c.llm_verdict:
                    lines.append(f"Verdict: {c.llm_verdict} — {c.llm_note or ''}")
                lines.append("")
            LightRAGService.get_instance().enqueue_ingest(
                "\n".join(lines),
                f"citation_run:{report.doc_name}:{report.generated_at}",
                "citation_verification",
            )
        except Exception:
            pass

        if emit_fn:
            s = report.summary
            await emit_fn("agent_update", {
                "agent_id": "citation_verifier",
                "status":   "done",
                "summary":  (
                    f"{report.doc_name}: {s['confirmed']} confirmed, "
                    f"{s.get('hallucinated', 0)} not found, "
                    f"{s.get('unreachable', 0)} unreachable"
                ),
            })
            await emit_fn("render_canvas", {
                "title": f"Citation Report — {report.doc_name}",
                "blocks": [{"type": "document", "data": {
                    "title": f"{report.doc_name} ({s['total']} citations)",
                    "content": _build_canvas_summary(report),
                }}],
            })

        return {
            "document":    report.doc_name,
            "run_id":      run_id,
            "citations":   [
                {
                    "index":           c.index,
                    "claim":           c.claim[:300],
                    "source_url":      c.source_url,
                    "status":          c.status,
                    "match_score":     round(c.match_score, 3),
                    "matched_passage": (c.matched_passage or "")[:200],
                    "llm_verdict":     c.llm_verdict,
                    "llm_note":        c.llm_note,
                }
                for c in report.citations
            ],
            "summary":     report.summary,
            "report_path": report.pdf_path,
        }
    except Exception as exc:
        logger.error("[citation_verifier] tool_handler error: %s", exc)
        if emit_fn:
            await emit_fn("agent_update", {
                "agent_id": "citation_verifier",
                "status":   "done",
                "summary":  f"Verification failed: {exc}",
            })
        return {"error": str(exc)}
