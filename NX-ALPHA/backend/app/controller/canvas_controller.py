"""
AURA NX-Alpha — Canvas Controller

Canvas export and vision endpoints.

ROUTES:
    POST /canvas/export — Serialize canvas blocks to markdown/pdf/docx/html/txt
    POST /canvas/image  — Receive a dropped image, send to vision model, stream AURA's response
"""

import asyncio
import io
import logging
import re

from fastapi import APIRouter, BackgroundTasks, HTTPException, UploadFile
from fastapi.responses import Response
from pydantic import BaseModel
from typing import List, Optional

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/canvas", tags=["canvas"])


class ExportRequest(BaseModel):
    blocks: List[dict]
    format: str = "pdf"
    title: str = "aura-output"


# ─────────────────────────────────────────────────────────────────────────────
# MARKDOWN RENDERER (existing logic, kept as-is for .md export)
# ─────────────────────────────────────────────────────────────────────────────

def _blocks_to_markdown(blocks: list[dict]) -> str:
    """Serialize canvas blocks to markdown text."""
    lines = []
    for block in blocks:
        d = block.get("data", block)
        block_type = block.get("type", "")

        if block_type == "heading":
            level = d.get("level", 1)
            lines.append(f"{'#' * level} {d.get('text', '')}\n")

        elif block_type in ("paragraph", "document", "document_file"):
            content = d.get("content") or d.get("text", "")
            if d.get("title") and block_type in ("document", "document_file"):
                lines.append(f"## {d['title']}\n")
            lines.append(f"{content}\n")

        elif block_type == "code":
            lang = d.get("language", "")
            content = d.get("content") or d.get("code", "")
            lines.append(f"```{lang}\n{content}\n```\n")

        elif block_type == "list":
            items = d.get("items", [])
            for item in items:
                lines.append(f"- {item}")
            lines.append("")

        elif block_type == "table":
            headers = d.get("headers", [])
            rows = d.get("rows", [])
            if headers:
                lines.append("| " + " | ".join(str(h) for h in headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for row in rows:
                    lines.append("| " + " | ".join(str(c) for c in row) + " |")
                lines.append("")

        elif block_type == "html":
            src = d.get("src", "")
            title = d.get("title", src)
            lines.append(f"[{title}]({src})\n")

        elif block_type in ("image", "image_generated"):
            url = d.get("url") or d.get("src", "")
            alt = d.get("alt", "image")
            lines.append(f"![{alt}]({url})\n")

        elif block_type == "metric_card":
            label = d.get("label", "")
            value = d.get("value", "")
            lines.append(f"**{label}:** {value}\n")

        elif block_type == "callout":
            content = d.get("content") or d.get("text", "")
            lines.append(f"> {content}\n")

        elif block_type == "chart":
            title = d.get("title", "Chart")
            lines.append(f"**{title}** *(chart — not exportable to markdown)*\n")

        elif block_type == "email":
            to = d.get("to", "")
            subject = d.get("subject", "")
            body = d.get("body") or d.get("content", "")
            lines.append(f"**To:** {to}")
            lines.append(f"**Subject:** {subject}\n")
            lines.append(f"{body}\n")

        elif block_type == "diagram":
            content = d.get("content") or d.get("code", "")
            lines.append(f"```mermaid\n{content}\n```\n")

        else:
            content = d.get("content") or d.get("text") or d.get("value", "")
            if content:
                lines.append(f"{content}\n")

        lines.append("")  # blank line between blocks

    return "\n".join(lines).strip()


# ─────────────────────────────────────────────────────────────────────────────
# HTML RENDERER (used for PDF + HTML export)
# ─────────────────────────────────────────────────────────────────────────────

_CSS = """
body {
    font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, sans-serif;
    max-width: 800px; margin: 0 auto; padding: 40px 20px;
    color: #1a1a2e; background: #ffffff; line-height: 1.7;
}
h1, h2, h3 { color: #9a7b2d; margin-top: 1.5em; margin-bottom: 0.5em; }
h1 { font-size: 1.8em; border-bottom: 2px solid #C9A84C; padding-bottom: 0.3em; }
h2 { font-size: 1.4em; }
h3 { font-size: 1.15em; }
p { margin: 0.6em 0; }
pre {
    background: #1e1e2e; color: #cdd6f4; padding: 16px; border-radius: 8px;
    overflow-x: auto; font-size: 0.9em; line-height: 1.5;
}
code { font-family: 'Cascadia Code', 'Fira Code', monospace; }
table {
    width: 100%; border-collapse: collapse; margin: 1em 0;
    font-size: 0.95em;
}
th {
    background: #C9A84C; color: #fff; padding: 10px 12px;
    text-align: left; font-weight: 600;
}
td { padding: 8px 12px; border-bottom: 1px solid #e0e0e0; }
tr:nth-child(even) td { background: #f8f6f0; }
ul, ol { padding-left: 1.5em; margin: 0.5em 0; }
li { margin: 0.3em 0; }
blockquote {
    border-left: 4px solid #C9A84C; margin: 1em 0; padding: 0.5em 1em;
    background: #fdf8ec; color: #5a4e2e;
}
.metric { display: inline-block; background: #f0ead6; border-radius: 8px;
    padding: 12px 20px; margin: 6px 4px; text-align: center; }
.metric .value { font-size: 1.6em; font-weight: 700; color: #9a7b2d; }
.metric .label { font-size: 0.85em; color: #666; }
.email-block { background: #f8f8f8; border: 1px solid #ddd; border-radius: 8px;
    padding: 16px; margin: 1em 0; }
.email-header { font-weight: 600; margin-bottom: 4px; }
img { max-width: 100%; height: auto; border-radius: 6px; margin: 1em 0; }
a { color: #9a7b2d; }
.footer { margin-top: 3em; padding-top: 1em; border-top: 1px solid #e0e0e0;
    color: #999; font-size: 0.8em; text-align: center; }
@media print {
    body { padding: 0; }
    pre, table, blockquote { page-break-inside: avoid; }
}
"""


def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _blocks_to_html(blocks: list[dict], title: str) -> str:
    """Render canvas blocks as a full styled HTML document."""
    parts = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        f"<meta charset='utf-8'><title>{_escape_html(title)}</title>",
        f"<style>{_CSS}</style>",
        "</head><body>",
    ]

    for block in blocks:
        d = block.get("data", block)
        bt = block.get("type", "")

        if bt == "heading":
            level = min(d.get("level", 1), 3)
            parts.append(f"<h{level}>{_escape_html(d.get('text', ''))}</h{level}>")

        elif bt in ("paragraph", "document", "document_file"):
            content = d.get("content") or d.get("text", "")
            if d.get("title") and bt in ("document", "document_file"):
                parts.append(f"<h2>{_escape_html(d['title'])}</h2>")
            # Convert markdown-style links to HTML links
            content_html = re.sub(
                r'\[([^\]]+)\]\(([^)]+)\)',
                r'<a href="\2">\1</a>',
                _escape_html(content),
            )
            # Convert **bold** to <strong>
            content_html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', content_html)
            parts.append(f"<p>{content_html}</p>")

        elif bt == "code":
            lang = d.get("language", "")
            content = d.get("content") or d.get("code", "")
            parts.append(f"<pre><code class='language-{_escape_html(lang)}'>{_escape_html(content)}</code></pre>")

        elif bt == "list":
            items = d.get("items", [])
            ordered = d.get("ordered", False)
            tag = "ol" if ordered else "ul"
            items_html = "".join(f"<li>{_escape_html(str(it))}</li>" for it in items)
            parts.append(f"<{tag}>{items_html}</{tag}>")

        elif bt == "table":
            headers = d.get("headers", [])
            rows = d.get("rows", [])
            if headers:
                th = "".join(f"<th>{_escape_html(str(h))}</th>" for h in headers)
                tr_rows = ""
                for row in rows:
                    td = "".join(f"<td>{_escape_html(str(c))}</td>" for c in row)
                    tr_rows += f"<tr>{td}</tr>"
                parts.append(f"<table><thead><tr>{th}</tr></thead><tbody>{tr_rows}</tbody></table>")

        elif bt == "html":
            src = d.get("src", "")
            link_title = d.get("title", src)
            parts.append(f'<p><a href="{_escape_html(src)}">{_escape_html(link_title)}</a></p>')

        elif bt in ("image", "image_generated"):
            url = d.get("url") or d.get("src", "")
            alt = d.get("alt", "image")
            parts.append(f'<img src="{_escape_html(url)}" alt="{_escape_html(alt)}">')

        elif bt == "metric_card":
            label = d.get("label", "")
            value = d.get("value", "")
            parts.append(
                f'<div class="metric">'
                f'<div class="value">{_escape_html(str(value))}</div>'
                f'<div class="label">{_escape_html(label)}</div></div>'
            )

        elif bt == "callout":
            content = d.get("content") or d.get("text", "")
            parts.append(f"<blockquote>{_escape_html(content)}</blockquote>")

        elif bt == "chart":
            chart_title = d.get("title", "Chart")
            parts.append(f"<p><em>{_escape_html(chart_title)} (interactive chart — not available in export)</em></p>")

        elif bt == "email":
            to_addr = d.get("to", "")
            subject = d.get("subject", "")
            body = d.get("body") or d.get("content", "")
            parts.append(
                f'<div class="email-block">'
                f'<div class="email-header">To: {_escape_html(to_addr)}</div>'
                f'<div class="email-header">Subject: {_escape_html(subject)}</div>'
                f'<p>{_escape_html(body)}</p></div>'
            )

        elif bt == "diagram":
            content = d.get("content") or d.get("code", "")
            parts.append(f"<pre><code class='language-mermaid'>{_escape_html(content)}</code></pre>")

        else:
            content = d.get("content") or d.get("text") or d.get("value", "")
            if content:
                parts.append(f"<p>{_escape_html(str(content))}</p>")

    parts.append('<div class="footer">Generated by AURA</div>')
    parts.append("</body></html>")
    return "\n".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX RENDERER
# ─────────────────────────────────────────────────────────────────────────────

def _blocks_to_docx(blocks: list[dict], title: str) -> io.BytesIO:
    """Render canvas blocks as a DOCX document. Returns BytesIO buffer."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x9A, 0x7B, 0x2D)

    for block in blocks:
        d = block.get("data", block)
        bt = block.get("type", "")

        if bt == "heading":
            level = min(d.get("level", 1), 3)
            h = doc.add_heading(d.get("text", ""), level=level)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x9A, 0x7B, 0x2D)

        elif bt in ("paragraph", "document", "document_file"):
            content = d.get("content") or d.get("text", "")
            if d.get("title") and bt in ("document", "document_file"):
                doc.add_heading(d["title"], level=2)
            doc.add_paragraph(content)

        elif bt == "code":
            content = d.get("content") or d.get("code", "")
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

        elif bt == "list":
            items = d.get("items", [])
            list_style = 'List Number' if d.get("ordered") else 'List Bullet'
            for item in items:
                doc.add_paragraph(str(item), style=list_style)

        elif bt == "table":
            headers = d.get("headers", [])
            rows = d.get("rows", [])
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = 'Light Grid Accent 1'
                # Header row
                for j, h in enumerate(headers):
                    table.rows[0].cells[j].text = str(h)
                # Data rows
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        if j < len(headers):
                            table.rows[i + 1].cells[j].text = str(cell)

        elif bt in ("image", "image_generated"):
            url = d.get("url") or d.get("src", "")
            alt = d.get("alt", "image")
            doc.add_paragraph(f"[Image: {alt}] {url}")

        elif bt == "metric_card":
            label = d.get("label", "")
            value = d.get("value", "")
            p = doc.add_paragraph()
            run = p.add_run(f"{value}")
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x9A, 0x7B, 0x2D)
            p.add_run(f"  {label}")

        elif bt == "callout":
            content = d.get("content") or d.get("text", "")
            p = doc.add_paragraph()
            p.style = doc.styles['Intense Quote'] if 'Intense Quote' in doc.styles else doc.styles['Normal']
            p.add_run(content)

        elif bt == "email":
            to_addr = d.get("to", "")
            subject = d.get("subject", "")
            body = d.get("body") or d.get("content", "")
            doc.add_paragraph(f"To: {to_addr}\nSubject: {subject}")
            doc.add_paragraph(body)

        elif bt == "html":
            src = d.get("src", "")
            link_title = d.get("title", src)
            doc.add_paragraph(f"{link_title}: {src}")

        else:
            content = d.get("content") or d.get("text") or d.get("value", "")
            if content:
                doc.add_paragraph(str(content))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────────────────
# PLAIN TEXT RENDERER
# ─────────────────────────────────────────────────────────────────────────────

def _blocks_to_txt(blocks: list[dict]) -> str:
    """Render canvas blocks as plain text (no markdown syntax)."""
    parts = []
    for block in blocks:
        d = block.get("data", block)
        bt = block.get("type", "")

        if bt == "heading":
            parts.append(d.get("text", "").upper())
            parts.append("")

        elif bt in ("paragraph", "document", "document_file"):
            content = d.get("content") or d.get("text", "")
            if d.get("title") and bt in ("document", "document_file"):
                parts.append(d["title"].upper())
            parts.append(content)
            parts.append("")

        elif bt == "code":
            content = d.get("content") or d.get("code", "")
            parts.append(content)
            parts.append("")

        elif bt == "list":
            items = d.get("items", [])
            for i, item in enumerate(items):
                if d.get("ordered"):
                    parts.append(f"  {i+1}. {item}")
                else:
                    parts.append(f"  - {item}")
            parts.append("")

        elif bt == "table":
            headers = d.get("headers", [])
            rows = d.get("rows", [])
            if headers:
                parts.append("  ".join(str(h) for h in headers))
                parts.append("-" * 40)
                for row in rows:
                    parts.append("  ".join(str(c) for c in row))
            parts.append("")

        elif bt == "callout":
            content = d.get("content") or d.get("text", "")
            parts.append(f"  {content}")
            parts.append("")

        elif bt == "email":
            to_addr = d.get("to", "")
            subject = d.get("subject", "")
            body = d.get("body") or d.get("content", "")
            parts.append(f"To: {to_addr}")
            parts.append(f"Subject: {subject}")
            parts.append(body)
            parts.append("")

        else:
            content = d.get("content") or d.get("text") or d.get("value", "")
            if content:
                parts.append(str(content))
                parts.append("")

    return "\n".join(parts).strip()


# ─────────────────────────────────────────────────────────────────────────────
# EXPORT ENDPOINT
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/export")
async def export_canvas(req: ExportRequest):
    """Serialize canvas blocks to the requested format and return as file download."""
    safe_title = req.title.replace(' ', '-').lower()

    if req.format == "pdf":
        html = _blocks_to_html(req.blocks, req.title)
        try:
            from playwright.async_api import async_playwright
            async with async_playwright() as p:
                browser = await p.chromium.launch()
                page = await browser.new_page()
                await page.set_content(html, wait_until="domcontentloaded")
                pdf_bytes = await page.pdf(
                    format="A4",
                    print_background=True,
                    margin={"top": "1in", "bottom": "1in", "left": "1in", "right": "1in"},
                )
                await browser.close()
            logger.info("[canvas_controller] Export: %d blocks → %s.pdf (%d bytes)",
                        len(req.blocks), safe_title, len(pdf_bytes))
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
            )
        except Exception as exc:
            logger.error("[canvas_controller] PDF export failed: %s — falling back to HTML", exc)
            # Fall back to HTML if Playwright fails
            return Response(
                content=html.encode("utf-8"),
                media_type="text/html",
                headers={"Content-Disposition": f'attachment; filename="{safe_title}.html"'},
            )

    elif req.format == "html":
        html = _blocks_to_html(req.blocks, req.title)
        logger.info("[canvas_controller] Export: %d blocks → %s.html (%d bytes)",
                    len(req.blocks), safe_title, len(html))
        return Response(
            content=html.encode("utf-8"),
            media_type="text/html",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.html"'},
        )

    elif req.format == "docx":
        try:
            buf = _blocks_to_docx(req.blocks, req.title)
            content = buf.getvalue()
            logger.info("[canvas_controller] Export: %d blocks → %s.docx (%d bytes)",
                        len(req.blocks), safe_title, len(content))
            return Response(
                content=content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
            )
        except ImportError:
            logger.error("[canvas_controller] DOCX export failed: python-docx not installed")
            return Response(content=b"DOCX export requires python-docx", status_code=500)

    elif req.format == "txt":
        txt = _blocks_to_txt(req.blocks)
        logger.info("[canvas_controller] Export: %d blocks → %s.txt (%d bytes)",
                    len(req.blocks), safe_title, len(txt))
        return Response(
            content=txt.encode("utf-8"),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.txt"'},
        )

    else:
        # Default: markdown
        markdown_content = _blocks_to_markdown(req.blocks)
        filename = f"{safe_title}.md"
        logger.info("[canvas_controller] Export: %d blocks → %s (%d bytes)",
                    len(req.blocks), filename, len(markdown_content))
        return Response(
            content=markdown_content.encode("utf-8"),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )


# ─────────────────────────────────────────────────────────────────────────────
# CANVAS IMAGE — Vision analysis via Ollama
# ─────────────────────────────────────────────────────────────────────────────

class CanvasImageRequest(BaseModel):
    image_data: str              # base64 data URI: "data:image/png;base64,..."
    filename: str = "image"
    thread_id: Optional[str] = None  # optional — falls back to backend session


@router.post("/image")
async def canvas_image(req: CanvasImageRequest, background_tasks: BackgroundTasks):
    """Receive an image dropped on canvas, send to vision model, stream AURA's response."""
    background_tasks.add_task(_process_canvas_image, req)
    return {"status": "processing"}


async def _process_canvas_image(req: CanvasImageRequest) -> None:
    """
    Background task: send image to vision model, emit response as AURA tokens.
    Routes to Interface Engine (Qwen3-VL) when mmproj is loaded, falls back to Ollama.
    """
    from app.controller.chat_controller import _emit, _get_session_thread_id
    from app.service.interface_engine import get_engine
    from app.service.ollama_service import get_ollama_service
    from app.service.memory_service import get_memory_service

    try:
        # Strip data URI prefix to get raw base64
        raw_b64 = req.image_data.split(",", 1)[1] if "," in req.image_data else req.image_data

        vision_prompt = (
            f"The user just shared an image on our shared canvas work surface "
            f"(filename: {req.filename}). "
            "Describe what you see and engage with it naturally — offer observations, "
            "ask a relevant question, or suggest what we could do with it. "
            "Be direct and specific about what's in the image."
        )

        await _emit("thinking", {"text": "Looking at your image..."})

        # Prefer Interface Engine (local, always-on, no VRAM spike)
        engine = get_engine()
        if engine and engine.has_vision:
            result = await engine.generate_vision(raw_b64, vision_prompt, max_tokens=512)
            response = result["text"]
            logger.info("[canvas_controller] Image handled by Interface Engine (%dms)", result.get("latency_ms", 0))
        else:
            # Fallback: Ollama workhorse
            ollama_svc = get_ollama_service()
            if ollama_svc is None:
                await _emit("error", {"message": "Vision unavailable — no vision model loaded."})
                return
            response = await ollama_svc.chat_with_image(prompt=vision_prompt, image_b64=raw_b64)
            logger.info("[canvas_controller] Image handled by Ollama workhorse (mmproj not active)")

        # Emit as a single complete message
        msg_id = __import__("uuid").uuid4().hex
        await _emit("token", {"text": response, "messageId": msg_id})
        await _emit("end", {"reason": "vision_complete"})

        # Persist to conversation memory so follow-up chat has context
        thread_id = req.thread_id or _get_session_thread_id()
        mem_svc = get_memory_service()
        if mem_svc:
            mem_svc._append_sliding_window(thread_id, "user", f"[Shared image: {req.filename}]")
            mem_svc._append_sliding_window(thread_id, "assistant", response)

    except Exception as e:
        logger.error("[canvas_controller] Vision processing failed: %s", e)
        await _emit("error", {"message": f"Vision processing failed: {str(e)}"})


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT IMPORT — PDF / DOCX
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/document")
async def canvas_document(
    file: UploadFile,
    background_tasks: BackgroundTasks,
):
    """
    Accept a PDF or DOCX upload, extract text, and ingest into personal memory.
    The Interface Engine can then answer questions about the document via memory_search.
    """
    from app.config import get_settings
    settings = get_settings()
    max_bytes = settings.max_upload_mb * 1024 * 1024

    # Size guard (read up to max + 1 byte to detect over-limit)
    content = await file.read(max_bytes + 1)
    if len(content) > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size is {settings.max_upload_mb} MB.",
        )

    filename = file.filename or "document"
    background_tasks.add_task(_process_canvas_document, content, filename)
    return {"status": "processing", "filename": filename}


async def _process_canvas_document(file_bytes: bytes, filename: str) -> None:
    """Background task: extract text from document and ingest into personal memory."""
    from app.controller.chat_controller import _emit
    from app.controller.data_controller import _ingest_personal_document

    try:
        from app.service.document_parser import extract_text
        await _emit("thinking", {"text": f"Reading {filename}..."})

        text = extract_text(file_bytes, filename)
        chunk_count = await _ingest_personal_document(
            content=text,
            doc_type="user_context",
            title=filename,
            tags=["imported_document"],
        )

        # Build canvas preview — LLM summary if engine available, else raw excerpt
        import uuid as _uuid
        await _emit("thinking", {"text": f"Summarizing {filename}..."})

        summary = None
        try:
            from app.service.interface_engine import get_engine
            engine = get_engine()
            if engine is not None:
                sample = text[:3000]
                result = await engine.generate(
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "You are a concise document analyst. "
                                "Summarize the following document in 3-5 sentences. "
                                "Focus on what the document is about and its key points. "
                                "Be direct — no preamble."
                            ),
                        },
                        {"role": "user", "content": sample},
                    ],
                    max_tokens=256,
                    temperature=0.3,
                )
                summary = result.get("text", "").strip()
        except Exception as sum_err:
            logger.warning("[canvas_controller] Summary generation failed: %s", sum_err)

        # Fallback to raw excerpt if LLM unavailable or failed
        if not summary:
            summary = text[:800].rsplit(" ", 1)[0] + " …" if len(text) > 800 else text

        await _emit("render_canvas", {
            "title": filename,
            "blocks": [
                {
                    "id":   _uuid.uuid4().hex,
                    "type": "heading",
                    "data": {"text": filename, "level": 2},
                },
                {
                    "id":   _uuid.uuid4().hex,
                    "type": "paragraph",
                    "data": {"text": summary},
                },
            ],
        })

        msg = (
            f"**{filename}** loaded — {chunk_count} section(s) in memory. "
            "Ask me anything about it."
        )
        msg_id = _uuid.uuid4().hex
        await _emit("token", {"text": msg, "messageId": msg_id})
        await _emit("end", {"reason": "document_ingested"})
        logger.info("[canvas_controller] Document '%s' ingested: %d chunks", filename, chunk_count)

    except ValueError as e:
        # Unsupported format
        await _emit("error", {"message": str(e)})
    except Exception as e:
        logger.error("[canvas_controller] Document ingestion failed for '%s': %s", filename, e)
        await _emit("error", {"message": f"Could not read '{filename}': {str(e)}"})
